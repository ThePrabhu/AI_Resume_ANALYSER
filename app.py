import streamlit as st
import fitz  # PyMuPDF
import docx
import re
import pandas as pd
import io
from docx import Document
import google.generativeai as genai

# Set your Gemini API key
genai.configure(api_key="AIzaSyBshemiB4n-Zq7S1nKOZctmTfovGSYqn9M")  # Replace with your actual key

# PDF extractor
def extract_text_from_pdf(file):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    return "\n".join(page.get_text() for page in doc)

# DOCX extractor
def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join(para.text for para in doc.paragraphs)

# Extract section
def extract_section(text, header):
    sections = text.split("###")
    for section in sections:
        if header.lower() in section.lower():
            return section.strip()
    return "Section not found."

# Generate DOCX
def generate_docx(text):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Match jobs from dataset
def match_jobs(resume_text, jobs_df):
    matched_jobs = []
    for _, row in jobs_df.iterrows():
        score = sum(1 for skill in row['skills'].split(",") if skill.strip().lower() in resume_text.lower())
        if score >= 2:
            matched_jobs.append((row['job_title'], score))
    return sorted(matched_jobs, key=lambda x: x[1], reverse=True)

# Streamlit page setup
st.set_page_config(page_title="AI Resume Optimizer", layout="wide")
st.title("AI Resume Analyzer & Optimizer (Gemini)")
st.markdown("Upload your resume and let Gemini AI analyze and rewrite it for a specific job role!")

# Upload resume
uploaded_file = st.file_uploader("Upload your Resume (PDF or DOCX)", type=["pdf", "docx"])
resume_text = ""

if uploaded_file:
    file_type = uploaded_file.name.split(".")[-1]
    if file_type == "pdf":
        resume_text = extract_text_from_pdf(uploaded_file)
    elif file_type == "docx":
        resume_text = extract_text_from_docx(uploaded_file)
    else:
        st.error("Unsupported file format.")

    # Display extracted resume
    st.subheader("Extracted Resume Text")
    st.text_area("Here’s what we extracted:", resume_text, height=250)

# Target role input
if resume_text:
    st.subheader("Enter Target Job Role")
    target_role = st.text_input("Example: Data Analyst, Backend Developer")

    if target_role and st.button("Analyze & Rewrite Resume"):
        with st.spinner("Gemini AI is analyzing your resume..."):
            try:
                model = genai.GenerativeModel("gemini-1.5-flash")

                prompt = f"""
You are a professional resume coach.

Here is the user's resume:
\"\"\"{resume_text}\"\"\"

The user is targeting the role of: {target_role}

Your tasks:
1. Review the resume and give feedback (tone, grammar, format, keyword match)
2. Rewrite the resume professionally for the role of '{target_role}'
3. Suggest at least 3 other suitable job titles this user can apply for
4. List 5 missing or trending skills they should add
5. Score their resume out of 10 based on quality & role match

Respond clearly using the following format:

### 1. Resume Review
...

### 2. Rewritten Resume
...

### 3. Job Title Suggestions
...

### 4. Skills to Add
...

### 5. Resume Score
...
"""

                response = model.generate_content(prompt)
                ai_output = response.text

                review = extract_section(ai_output, "Review")
                rewritten = extract_section(ai_output, "Rewritten Resume")
                skills = extract_section(ai_output, "Skills to Add")
                job_titles = extract_section(ai_output, "Job Title Suggestions")
                score = extract_section(ai_output, "Resume Score")

            except Exception as e:
                st.error(f"Error: {e}")
                st.stop()

        # Resume comparison
        st.markdown("---")
        st.subheader("Resume Comparison")

        col1, col2 = st.columns(2)
        with col1:
            st.markdown("Original Resume")
            st.text_area("Original", resume_text, height=400)

        with col2:
            st.markdown("AI-Rewritten Resume")
            st.text_area("Improved", rewritten, height=400)

        with st.expander("Gemini's Resume Review"):
            st.markdown(review)

        with st.expander("Skills You Should Add"):
            st.markdown(skills)

        with st.expander("Suggested Job Titles for You"):
            st.markdown(job_titles)

        with st.expander("Resume Score"):
            st.markdown(score)

        score_match = re.search(r"(\d{1,2})/10", score)
        if score_match:
            numeric_score = int(score_match.group(1))
            st.subheader("Resume Score Progress")
            st.progress(numeric_score / 10)

        # Download TXT
        st.download_button(
            label="Download Rewritten Resume (TXT)",
            data=rewritten,
            file_name="optimized_resume.txt",
            mime="text/plain"
        )

        # Download DOCX
        docx_data = generate_docx(rewritten)
        st.download_button("Download Rewritten Resume (DOCX)", data=docx_data, file_name="optimized_resume.docx")

        # Optional Static Job Matching (needs jobs.csv)
        try:
            jobs = pd.read_csv("jobs.csv")
            job_matches = match_jobs(resume_text, jobs)
            if job_matches:
                with st.expander("Matching Job Titles Based on Your Resume"):
                    for title, score in job_matches[:5]:
                        st.markdown(f"{title} — Skill Match Score: {score}")
        except Exception:
            st.warning("jobs.csv not found or could not be loaded. Upload a job dataset to enable matching.")